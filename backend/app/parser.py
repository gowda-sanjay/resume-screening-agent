import io
import fitz  # PyMuPDF
import docx

def parse_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file using python-docx.
    """
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file using PyMuPDF (fitz).
    """
    text = ""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"
        doc.close()
        return text
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")

def parse_document(file_name: str, file_bytes: bytes) -> str:
    """
    Delegates to appropriate parser based on file extension.
    """
    ext = file_name.split('.')[-1].lower()
    if ext == 'pdf':
        return parse_pdf(file_bytes)
    elif ext in ['docx', 'doc']:
        return parse_docx(file_bytes)
    elif ext == 'txt':
        return file_bytes.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file format: .{ext}. Must be PDF, DOCX, or TXT.")
